"""
t_document_generator.py

Business-level legal document generation engine — India jurisdiction.

This file contains ONLY the core generation logic.
Prompts   → t_prompts.py
Intent    → t_intent.py
DB        → s_db.py
Routes    → t_document_route.py
"""

import os
import re
import json
import logging
from io import BytesIO
from typing import Any

from openai import AsyncOpenAI

# Import prompts and intent from separate modules
from feature_modules.prompts import EXTRACTION_PROMPTS, GENERATION_PROMPTS
from feature_modules.intent  import classify_intent, resolve_document_type

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Model config
# ---------------------------------------------------------------------------

_MODEL   = os.environ.get("MODEL_NAME", "gpt-5-nano")
_API_KEY = os.environ.get("OPENAI_API_KEY", "")

_NO_TEMP = {"gpt-5-nano", "gpt-4.1-nano", "gpt-4o-mini", "o1", "o1-mini", "o3-mini", "o3"}
_MCT     = {"gpt-5-nano", "gpt-4.1-nano", "gpt-4o-mini", "o1", "o1-mini", "o3-mini", "o3"}


def _get_client() -> AsyncOpenAI:
    key = os.environ.get("OPENAI_API_KEY", _API_KEY)
    return AsyncOpenAI(api_key=key)


def _api_kwargs(max_tokens: int, use_json: bool = False) -> dict:
    """
    Build OpenAI API kwargs.
    use_json=True  → field extraction   (response_format: json_object)
    use_json=False → document generation (plain text, NO response_format)

    IMPORTANT: Never set response_format for document generation —
    gpt-5-nano returns empty content when json_object is set on plain-text calls.
    """
    model  = os.environ.get("MODEL_NAME", _MODEL)
    kwargs: dict = {"model": model}
    if model not in _NO_TEMP:
        kwargs["temperature"] = 0.1
    if model in _MCT:
        kwargs["max_completion_tokens"] = max_tokens
    else:
        kwargs["max_tokens"] = max_tokens
    if use_json:
        kwargs["response_format"] = {"type": "json_object"}
    return kwargs


# ---------------------------------------------------------------------------
# Supported document types
# ---------------------------------------------------------------------------

SUPPORTED_DOCUMENT_TYPES: dict[str, str] = {
    "nda":                   "Non-Disclosure Agreement",
    "job_offer":             "Job Offer Letter",
    "freelancer_agreement":  "Freelancer Agreement",
    "service_agreement":     "Service Agreement",
    "consulting_agreement":  "Consulting Agreement",
    "lease_agreement":       "Lease Agreement",
    "employment_contract":   "Employment Contract",
}

# ---------------------------------------------------------------------------
# Field schemas
# ---------------------------------------------------------------------------

_SCHEMAS: dict[str, dict[str, list[str]]] = {
    "nda": {
        "required": ["disclosing_party", "receiving_party", "effective_date",
                     "purpose", "duration_years", "governing_state"],
        "optional": ["confidential_info_description", "exclusions",
                     "return_of_materials", "remedies", "signatory_names",
                     "stamp_duty_state"],
    },
    "job_offer": {
        "required": ["candidate_name", "company_name", "job_title",
                     "start_date", "ctc", "employment_type", "reporting_to"],
        "optional": ["work_location", "probation_period", "benefits",
                     "esop", "joining_bonus", "offer_expiry_date",
                     "hr_contact", "notice_period"],
    },
    "freelancer_agreement": {
        "required": ["client_name", "freelancer_name", "project_description",
                     "start_date", "end_date", "payment_amount",
                     "payment_schedule", "governing_state"],
        "optional": ["deliverables", "revision_rounds", "intellectual_property",
                     "confidentiality", "kill_fee", "late_payment_penalty",
                     "gst_applicable"],
    },
    "service_agreement": {
        "required": ["service_provider", "client", "services_description",
                     "start_date", "end_date", "fee", "payment_terms",
                     "governing_state"],
        "optional": ["service_levels", "termination_notice",
                     "limitation_of_liability", "indemnification",
                     "insurance_requirements", "dispute_resolution",
                     "gst_number"],
    },
    "consulting_agreement": {
        "required": ["consultant_name", "client_name", "scope_of_work",
                     "start_date", "end_date", "consulting_fee",
                     "payment_terms", "governing_state"],
        "optional": ["expenses_reimbursement", "non_compete_period",
                     "non_solicitation", "intellectual_property",
                     "confidentiality_period", "termination_clause",
                     "tds_applicable"],
    },
    "lease_agreement": {
        "required": ["landlord_name", "tenant_name", "property_address",
                     "lease_start_date", "lease_end_date", "monthly_rent",
                     "security_deposit", "governing_state"],
        "optional": ["maintenance_charges", "pet_policy",
                     "utilities_included", "lock_in_period",
                     "subletting_policy", "renewal_terms",
                     "notice_period", "stamp_duty_value"],
    },
    "employment_contract": {
        "required": ["employer_name", "employee_name", "job_title",
                     "department", "start_date", "ctc",
                     "work_hours", "governing_state"],
        "optional": ["probation_period", "notice_period", "non_compete",
                     "non_solicitation", "benefits", "leave_policy",
                     "termination_clause", "intellectual_property_assignment",
                     "pf_applicable", "gratuity_applicable"],
    },
}


# ---------------------------------------------------------------------------
# Own JSON extractor — no dependency on s_json_utils
# ---------------------------------------------------------------------------

def _parse_json(raw: str) -> dict:
    text = raw.replace("```json", "").replace("```", "").strip()
    try:
        data = json.loads(text)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            data = json.loads(match.group())
            if isinstance(data, dict):
                return data
        except json.JSONDecodeError:
            pass
    logger.warning("[doc_gen] could not parse JSON from model response")
    return {}


# ---------------------------------------------------------------------------
# Step 1: Extract fields using EXTRACTION_PROMPTS from t_prompts.py
# ---------------------------------------------------------------------------

async def _extract_fields(document_type: str, user_query: str) -> dict[str, Any]:
    schema = _SCHEMAS[document_type]
    system = EXTRACTION_PROMPTS[document_type]   # ← from t_prompts.py

    client = _get_client()
    # use_json=True → sets response_format=json_object for reliable JSON output
    kwargs = _api_kwargs(max_tokens=4096, use_json=True)
    kwargs["messages"] = [
        {"role": "system", "content": system},
        {"role": "user",   "content": (
            f"Extract all fields from this description as a JSON object.\n\n"
            f"Description:\n\"\"\"{user_query}\"\"\""
        )},
    ]

    response = await client.chat.completions.create(**kwargs)
    raw      = response.choices[0].message.content or "{}"

    logger.info(
        f"[doc_gen] field extraction "
        f"tokens={response.usage.prompt_tokens}in/{response.usage.completion_tokens}out "
        f"finish={response.choices[0].finish_reason}"
    )
    logger.debug(f"[doc_gen] extraction raw: {raw[:300]}")

    fields = _parse_json(raw)

    # Ensure all required fields exist
    for field in schema["required"]:
        if field not in fields or not fields[field]:
            fields[field] = "Not Specified"

    found = {k: v for k, v in fields.items() if v != "Not Specified"}
    logger.info(f"[doc_gen] extracted {len(found)} non-empty fields: {list(found.keys())}")

    return fields


# ---------------------------------------------------------------------------
# Step 2: Validate required fields
# ---------------------------------------------------------------------------

def _get_missing_fields(document_type: str, fields: dict) -> list[str]:
    return [
        f for f in _SCHEMAS[document_type]["required"]
        if not fields.get(f) or fields.get(f) == "Not Specified"
    ]


# ---------------------------------------------------------------------------
# Step 3: Generate document text using GENERATION_PROMPTS from t_prompts.py
# ---------------------------------------------------------------------------

async def _generate_document_text(
    document_type: str,
    fields: dict[str, Any],
    user_query: str,
) -> tuple[str, int, int]:
    doc_name   = SUPPORTED_DOCUMENT_TYPES[document_type]
    gen_prompt = GENERATION_PROMPTS[document_type]   # ← from t_prompts.py

    field_lines = "\n".join(
        f"  {k.replace('_', ' ').title()}: {v}"
        for k, v in fields.items()
        if v and v != "Not Specified"
    )

    system = gen_prompt

    user = f"""Use these details to generate the complete {doc_name} under Indian law:

{field_lines if field_lines else "(Use standard Indian law template values)"}

Additional context from user:
\"\"\"{user_query}\"\"\"

Write the complete {doc_name} now. Start directly with the document title."""

    client = _get_client()
    # CRITICAL: use_json=False — never set response_format for plain text generation
    kwargs = _api_kwargs(max_tokens=32000, use_json=False)
    kwargs["messages"] = [
        {"role": "system", "content": system},
        {"role": "user",   "content": user},
    ]

    response      = await client.chat.completions.create(**kwargs)
    document_text = response.choices[0].message.content or ""
    in_tok        = response.usage.prompt_tokens
    out_tok       = response.usage.completion_tokens
    finish        = response.choices[0].finish_reason

    logger.info(
        f"[doc_gen] document generated — type={document_type} "
        f"length={len(document_text)} chars "
        f"tokens={in_tok}in/{out_tok}out finish={finish}"
    )

    if not document_text.strip():
        logger.error(
            f"[doc_gen] empty response — finish={finish} "
            f"model={os.environ.get('MODEL_NAME', _MODEL)}"
        )

    return document_text, in_tok, out_tok


# ---------------------------------------------------------------------------
# Step 4: Render DOCX
# ---------------------------------------------------------------------------

def _render_docx(document_text: str, document_name: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime

    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    BLUE  = RGBColor(0x2c, 0x4a, 0x8c)
    DARK  = RGBColor(0x1a, 0x1a, 0x1a)
    MUTED = RGBColor(0x55, 0x55, 0x55)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    def _set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _remove_table_borders(table):
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tblBorders.append(b)
        tblPr.append(tblBorders)

    # Header banner
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell  = table.cell(0, 0)
    _set_cell_bg(cell, "1a2744")
    _remove_table_borders(table)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run(document_name.upper())
    run.bold           = True
    run.font.size      = Pt(20)
    run.font.color.rgb = WHITE

    para2 = cell.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2  = para2.add_run("Governed by the Laws of India")
    run2.font.size      = Pt(9)
    run2.font.color.rgb = RGBColor(0xCC, 0xD6, 0xF0)

    doc.add_paragraph()

    dp  = doc.add_paragraph()
    dr  = dp.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    dr.font.size      = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    dp.paragraph_format.space_after = Pt(12)

    lines        = document_text.split("\n")
    in_signature = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(3)
            continue

        if re.match(r'^\d+\.\s+SIGNATURE', stripped.upper()):
            in_signature = True

        is_section = (
            re.match(r'^\d+\.\s+[A-Z][A-Z\s,/&().-]{3,}[:.]?\s*$', stripped)
            or re.match(r'^[A-Z][A-Z\s,/&().-]{3,}[:.]?\s*$', stripped)
            or (stripped.isupper() and 4 < len(stripped) < 80)
        )

        if is_section and not in_signature:
            p    = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top  = OxmlElement("w:top")
            top.set(qn("w:val"),   "single")
            top.set(qn("w:sz"),    "4")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), "CCCCCC")
            pBdr.append(top)
            pPr.append(pBdr)
            run = p.add_run(stripped)
            run.bold           = True
            run.font.size      = Pt(11)
            run.font.color.rgb = BLUE
            continue

        if re.match(r'^\d+\.\d+\.?\s+', stripped) and not in_signature:
            p   = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.font.size      = Pt(10)
            run.font.color.rgb = DARK
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if in_signature and re.match(
            r'^(Name|Title|Date|Company|Designation|Signature|Witness|Place|PAN|Aadhaar)\s*:',
            stripped, re.I
        ):
            label, _, rest = stripped.partition(":")
            lp  = doc.add_paragraph()
            lr  = lp.add_run(label.strip() + ":")
            lr.bold            = True
            lr.font.size       = Pt(9)
            lr.font.color.rgb  = MUTED
            lp.paragraph_format.space_after = Pt(2)
            vp  = doc.add_paragraph()
            vr  = vp.add_run(rest.strip() if rest.strip() else "________________________")
            vr.font.size      = Pt(10)
            vr.font.color.rgb = DARK
            vp.paragraph_format.space_after = Pt(12)
            continue

        p   = doc.add_paragraph(style="Normal")
        run = p.add_run(stripped)
        run.font.size      = Pt(10)
        run.font.color.rgb = DARK
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)

    # Footer
    for section in doc.sections:
        footer      = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"{document_name}  |  Governed by Laws of India  |  Confidential  |  Page "
        )
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        fldChar1              = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText             = OxmlElement("w:instrText")
        instrText.text        = "PAGE"
        fldChar2              = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run2 = footer_para.add_run()
        run2.font.size      = Pt(8)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)

    buffer     = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()

    logger.info(f"[doc_gen] DOCX rendered — {len(docx_bytes):,} bytes")
    return docx_bytes


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def generate_document(
    document_type: str | None,
    user_query: str,
) -> dict:
    """
    Main entry point.

    document_type: canonical slug OR None (auto-detected via t_intent.classify_intent)
    user_query   : free-text description

    Returns dict with keys:
      status, document_type, document_name, fields, missing_fields,
      document, docx_bytes, word_count, input_tokens, output_tokens
    """
    total_in_tok  = 0
    total_out_tok = 0

    # ── Resolve or auto-detect document type ──────────────────────────────
    if document_type:
        doc_type = resolve_document_type(document_type)   # ← from t_intent.py
        if not doc_type:
            return {
                "status":         "unknown_type",
                "document_type":  document_type,
                "document_name":  "",
                "fields":         {},
                "missing_fields": [],
                "document":       "",
                "docx_bytes":     b"",
                "word_count":     0,
                "input_tokens":   0,
                "output_tokens":  0,
                "message": (
                    f"'{document_type}' is not supported. "
                    f"Supported: {', '.join(SUPPORTED_DOCUMENT_TYPES.keys())}"
                ),
            }
    else:
        logger.info("[doc_gen] document_type not provided — classifying intent via t_intent...")
        doc_type = await classify_intent(user_query)      # ← from t_intent.py
        if not doc_type:
            return {
                "status":         "unknown_type",
                "document_type":  None,
                "document_name":  "",
                "fields":         {},
                "missing_fields": [],
                "document":       "",
                "docx_bytes":     b"",
                "word_count":     0,
                "input_tokens":   total_in_tok,
                "output_tokens":  total_out_tok,
                "message": (
                    "Could not determine document type from description. "
                    "Please specify document_type explicitly. "
                    f"Supported: {', '.join(SUPPORTED_DOCUMENT_TYPES.keys())}"
                ),
            }

    doc_name = SUPPORTED_DOCUMENT_TYPES[doc_type]
    logger.info(f"[doc_gen] generating '{doc_name}' ({len(user_query)} chars)")

    # ── Step 1: Extract fields ─────────────────────────────────────────────
    fields = await _extract_fields(doc_type, user_query)

    # ── Step 2: Check missing ──────────────────────────────────────────────
    missing = _get_missing_fields(doc_type, fields)
    if missing:
        logger.warning(f"[doc_gen] missing required fields: {missing}")

    # ── Step 3: Generate document text ────────────────────────────────────
    document_text, in_tok, out_tok = await _generate_document_text(
        doc_type, fields, user_query
    )
    total_in_tok  += in_tok
    total_out_tok += out_tok

    # ── Step 4: Render DOCX ───────────────────────────────────────────────
    docx_bytes = _render_docx(document_text, doc_name)

    return {
        "status":         "missing_fields" if missing else "success",
        "document_type":  doc_type,
        "document_name":  doc_name,
        "fields":         fields,
        "missing_fields": missing,
        "document":       document_text,
        "docx_bytes":     docx_bytes,
        "word_count":     len(document_text.split()),
        "input_tokens":   total_in_tok,
        "output_tokens":  total_out_tok,
    }
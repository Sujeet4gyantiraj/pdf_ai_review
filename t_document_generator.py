"""
t_document_generator.py

Business-level legal document generation engine — INDIA JURISDICTION.

All documents are drafted under Indian law:
  - Indian Contract Act, 1872
  - Specific Relief Act, 1963
  - Information Technology Act, 2000
  - Indian Stamp Act, 1899
  - Transfer of Property Act, 1882 (lease)
  - Industrial Disputes Act, 1947 (employment)
  - Copyright Act, 1957 (IP)

Supported types:
  nda, job_offer, freelancer_agreement, service_agreement,
  consulting_agreement, lease_agreement, employment_contract
"""

import os
import re
import json
import logging
from io import BytesIO
from typing import Any

from openai import AsyncOpenAI

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Model config
# ---------------------------------------------------------------------------

_MODEL   = os.environ.get("MODEL_NAME", "gpt-5-nano")
_API_KEY = os.environ.get("OPENAI_API_KEY", "")

_NO_TEMP = {"gpt-5-nano", "gpt-4.1-nano", "gpt-4o-mini", "o1", "o1-mini", "o3-mini", "o3"}
_MCT     = {"gpt-5-nano", "gpt-4.1-nano", "gpt-4o-mini", "o1", "o1-mini", "o3-mini", "o3"}


def _get_client() -> AsyncOpenAI:
    key = _API_KEY or os.environ.get("OPENAI_API_KEY", "")
    return AsyncOpenAI(api_key=key)


def _api_kwargs(max_tokens: int, use_json: bool = False) -> dict:
    """
    Build OpenAI API kwargs.
    use_json=True  → field extraction   (response_format: json_object)
    use_json=False → document generation (plain text, NO response_format)
    """
    model  = os.environ.get("MODEL_NAME", _MODEL)
    kwargs: dict = {"model": model}
    if model not in _NO_TEMP:
        kwargs["temperature"] = 0.1
    if model in _MCT:
        kwargs["max_completion_tokens"] = max_tokens
    else:
        kwargs["max_tokens"] = max_tokens
    if use_json:
        kwargs["response_format"] = {"type": "json_object"}
    return kwargs


# ---------------------------------------------------------------------------
# Supported document types
# ---------------------------------------------------------------------------

SUPPORTED_DOCUMENT_TYPES: dict[str, str] = {
    "nda":                   "Non-Disclosure Agreement",
    "job_offer":             "Job Offer Letter",
    "freelancer_agreement":  "Freelancer Agreement",
    "service_agreement":     "Service Agreement",
    "consulting_agreement":  "Consulting Agreement",
    "lease_agreement":       "Lease Agreement",
    "employment_contract":   "Employment Contract",
}

_TYPE_ALIASES: dict[str, str] = {
    "non disclosure agreement":  "nda",
    "non-disclosure agreement":  "nda",
    "non_disclosure_agreement":  "nda",
    "confidentiality agreement": "nda",
    "job offer":                 "job_offer",
    "offer letter":              "job_offer",
    "job offer letter":          "job_offer",
    "freelance agreement":       "freelancer_agreement",
    "freelancer":                "freelancer_agreement",
    "service":                   "service_agreement",
    "consulting":                "consulting_agreement",
    "consultant agreement":      "consulting_agreement",
    "lease":                     "lease_agreement",
    "rental agreement":          "lease_agreement",
    "leave and licence":         "lease_agreement",
    "employment":                "employment_contract",
    "employment agreement":      "employment_contract",
    "appointment letter":        "job_offer",
}


def resolve_document_type(raw: str) -> str | None:
    cleaned  = raw.lower().strip().replace("-", "_").replace(" ", "_")
    if cleaned in SUPPORTED_DOCUMENT_TYPES:
        return cleaned
    readable = raw.lower().strip()
    if readable in _TYPE_ALIASES:
        return _TYPE_ALIASES[readable]
    for alias, slug in _TYPE_ALIASES.items():
        if alias in readable or readable in alias:
            return slug
    for slug in SUPPORTED_DOCUMENT_TYPES:
        if slug in cleaned or cleaned in slug:
            return slug
    return None


# ---------------------------------------------------------------------------
# Field schemas
# ---------------------------------------------------------------------------

_SCHEMAS: dict[str, dict[str, list[str]]] = {
    "nda": {
        "required": ["disclosing_party", "receiving_party", "effective_date",
                     "purpose", "duration_years", "governing_state"],
        "optional": ["confidential_info_description", "exclusions",
                     "return_of_materials", "remedies", "signatory_names",
                     "stamp_duty_state"],
    },
    "job_offer": {
        "required": ["candidate_name", "company_name", "job_title",
                     "start_date", "ctc", "employment_type", "reporting_to"],
        "optional": ["work_location", "probation_period", "benefits",
                     "esop", "joining_bonus", "offer_expiry_date",
                     "hr_contact", "notice_period"],
    },
    "freelancer_agreement": {
        "required": ["client_name", "freelancer_name", "project_description",
                     "start_date", "end_date", "payment_amount",
                     "payment_schedule", "governing_state"],
        "optional": ["deliverables", "revision_rounds", "intellectual_property",
                     "confidentiality", "kill_fee", "late_payment_penalty",
                     "gst_applicable"],
    },
    "service_agreement": {
        "required": ["service_provider", "client", "services_description",
                     "start_date", "end_date", "fee", "payment_terms",
                     "governing_state"],
        "optional": ["service_levels", "termination_notice",
                     "limitation_of_liability", "indemnification",
                     "insurance_requirements", "dispute_resolution",
                     "gst_number"],
    },
    "consulting_agreement": {
        "required": ["consultant_name", "client_name", "scope_of_work",
                     "start_date", "end_date", "consulting_fee",
                     "payment_terms", "governing_state"],
        "optional": ["expenses_reimbursement", "non_compete_period",
                     "non_solicitation", "intellectual_property",
                     "confidentiality_period", "termination_clause",
                     "tds_applicable"],
    },
    "lease_agreement": {
        "required": ["landlord_name", "tenant_name", "property_address",
                     "lease_start_date", "lease_end_date", "monthly_rent",
                     "security_deposit", "governing_state"],
        "optional": ["maintenance_charges", "pet_policy",
                     "utilities_included", "lock_in_period",
                     "subletting_policy", "renewal_terms",
                     "notice_period", "stamp_duty_value"],
    },
    "employment_contract": {
        "required": ["employer_name", "employee_name", "job_title",
                     "department", "start_date", "ctc",
                     "work_hours", "governing_state"],
        "optional": ["probation_period", "notice_period", "non_compete",
                     "non_solicitation", "benefits", "leave_policy",
                     "termination_clause", "intellectual_property_assignment",
                     "pf_applicable", "gratuity_applicable"],
    },
}


# ---------------------------------------------------------------------------
# Per-document-type EXTRACTION prompts — India specific fields
# ---------------------------------------------------------------------------

_EXTRACTION_PROMPTS: dict[str, str] = {

    "nda": """You are a legal assistant specialising in Indian Non-Disclosure Agreements.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
disclosing_party    - name of company/person sharing confidential info
receiving_party     - name of company/person receiving confidential info
effective_date      - when the NDA takes effect (DD/MM/YYYY or Month DD, YYYY)
purpose             - why confidential information is being shared
duration_years      - how long the NDA lasts (e.g. "2 years")
governing_state     - Indian state whose laws govern (e.g. "Maharashtra", "Karnataka")

Optional JSON keys:
confidential_info_description - what specific info is covered
exclusions          - what is NOT confidential
return_of_materials - must materials be returned on termination
remedies            - remedies for breach
signatory_names     - names of signatories
stamp_duty_state    - state for stamp duty purposes

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",

    "job_offer": """You are an HR specialist drafting Indian Job Offer / Appointment Letters.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
candidate_name  - full name of the candidate
company_name    - name of the company
job_title       - position being offered
start_date      - proposed date of joining (DD/MM/YYYY)
ctc             - Cost to Company (annual CTC in INR, e.g. "12,00,000 INR per annum")
employment_type - full-time / part-time / contract / internship
reporting_to    - manager or supervisor name/designation

Optional JSON keys:
work_location   - city and office address or remote
probation_period - probation duration (e.g. "6 months")
benefits        - PF, gratuity, health insurance, etc.
esop            - ESOP or stock options
joining_bonus   - one-time joining bonus in INR
offer_expiry_date - deadline to accept offer
hr_contact      - HR contact name and email
notice_period   - notice period after probation

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",

    "freelancer_agreement": """You are a contracts specialist drafting Indian Freelancer Agreements.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
client_name          - name of company/person hiring the freelancer
freelancer_name      - name of the freelancer
project_description  - what work is being done
start_date           - when work begins
end_date             - when work is due
payment_amount       - total payment in INR (e.g. "50,000 INR")
payment_schedule     - when/how payment is made (milestone/on-completion/monthly)
governing_state      - Indian state governing the agreement

Optional JSON keys:
deliverables         - specific outputs expected
revision_rounds      - number of revision rounds included
intellectual_property - who owns the work product
confidentiality      - confidentiality requirements
kill_fee             - cancellation fee
late_payment_penalty - late payment penalty
gst_applicable       - whether GST is applicable and GST numbers

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",

    "service_agreement": """You are a contracts specialist drafting Indian Service Agreements.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
service_provider     - name of company/person providing services
client               - name of company/person receiving services
services_description - description of the services
start_date           - service start date
end_date             - service end date
fee                  - total fee or rate in INR
payment_terms        - payment schedule (Net 30, monthly advance, etc.)
governing_state      - Indian state governing the agreement

Optional JSON keys:
service_levels       - SLA or performance standards
termination_notice   - notice period to terminate
limitation_of_liability - liability cap amount
indemnification      - indemnification terms
insurance_requirements - required insurance coverage
dispute_resolution   - arbitration or court
gst_number           - GST registration numbers of both parties

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",

    "consulting_agreement": """You are a contracts specialist drafting Indian Consulting Agreements.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
consultant_name  - name of consultant or firm
client_name      - name of client company
scope_of_work    - description of consulting services
start_date       - engagement start date
end_date         - engagement end date
consulting_fee   - fee in INR (hourly/daily/project basis)
payment_terms    - how and when payment is made
governing_state  - Indian state governing the agreement

Optional JSON keys:
expenses_reimbursement - expense reimbursement policy
non_compete_period     - non-compete duration after engagement
non_solicitation       - non-solicitation clause
intellectual_property  - ownership of work product
confidentiality_period - how long confidentiality lasts
termination_clause     - termination conditions and notice
tds_applicable         - whether TDS is deductible and at what rate

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",

    "lease_agreement": """You are a real estate attorney drafting Indian Lease / Leave and Licence Agreements.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
landlord_name    - full name of landlord/licensor
tenant_name      - full name of tenant(s)/licensee(s)
property_address - full address of the property including city, state, PIN code
lease_start_date - when the lease/licence begins
lease_end_date   - when the lease/licence ends
monthly_rent     - monthly rent in INR
security_deposit - security deposit in INR (usually 2-6 months rent)
governing_state  - Indian state where property is located

Optional JSON keys:
maintenance_charges - monthly maintenance amount
pet_policy          - whether pets are allowed
utilities_included  - electricity, water, gas included or not
lock_in_period      - minimum lock-in period
subletting_policy   - whether subletting is allowed
renewal_terms       - how the lease can be renewed
notice_period       - notice required to vacate
stamp_duty_value    - stamp duty value for registration

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",

    "employment_contract": """You are an employment law specialist drafting Indian Employment Contracts.
Extract fields from the user description. Return ONLY a valid JSON object.

Required JSON keys:
employer_name  - name of the employing company
employee_name  - full name of employee
job_title      - employee job title / designation
department     - department or team
start_date     - date of joining (DD/MM/YYYY)
ctc            - annual Cost to Company in INR
work_hours     - hours per day or week (e.g. "9 hours/day, 5 days/week")
governing_state - Indian state whose laws govern the contract

Optional JSON keys:
probation_period   - probationary period duration (typically 3-6 months)
notice_period      - notice period required by both parties
non_compete        - non-compete restrictions after leaving
non_solicitation   - non-solicitation clause
benefits           - PF, ESI, health insurance, gratuity, etc.
leave_policy       - casual leave, sick leave, earned leave entitlements
termination_clause - grounds and process for termination
intellectual_property_assignment - IP ownership clause
pf_applicable      - whether PF/EPF is applicable
gratuity_applicable - whether gratuity is applicable

Set any field not mentioned to "Not Specified". Return ONLY the JSON object.""",
}


# ---------------------------------------------------------------------------
# Per-document-type GENERATION prompts — India law specific
# ---------------------------------------------------------------------------

_GENERATION_PROMPTS: dict[str, str] = {

    "nda": """You are a senior advocate specialising in Indian contract and IP law.
Draft a complete, enforceable Non-Disclosure Agreement governed by Indian law.

Legal framework references:
- Indian Contract Act, 1872 (Sections 27, 73, 74)
- Information Technology Act, 2000
- Indian Stamp Act, 1899
- Copyright Act, 1957

Include these sections:
1. PARTIES AND RECITALS
2. DEFINITIONS
   2.1 Confidential Information
   2.2 Exclusions from Confidential Information
3. OBLIGATIONS OF RECEIVING PARTY
   3.1 Non-Disclosure Obligation
   3.2 Standard of Care
   3.3 Permitted Disclosures (including disclosures required by SEBI, RBI, or courts)
4. TERM AND TERMINATION
5. RETURN OR DESTRUCTION OF MATERIALS
6. REMEDIES AND RELIEF
   6.1 Injunctive Relief under Specific Relief Act, 1963
   6.2 Damages under Indian Contract Act, 1872
7. GENERAL PROVISIONS
   7.1 Governing Law and Jurisdiction (Indian courts)
   7.2 Dispute Resolution (arbitration under Arbitration and Conciliation Act, 1996)
   7.3 Entire Agreement
   7.4 Amendments
   7.5 Stamp Duty
8. SIGNATURE BLOCK

Formatting:
- Section headings in ALL CAPS with number
- Sub-clauses numbered 1.1, 1.2 etc.
- Formal legal language compliant with Indian law
- Use Indian currency (INR/Rupees) where applicable
- Reference relevant Indian statutes where appropriate
- Plain text only, no markdown
- Use actual values provided, use [TO BE AGREED] for missing values""",

    "job_offer": """You are a senior HR professional and employment law expert in India.
Draft a complete Job Offer / Appointment Letter compliant with Indian labour law.

Legal framework references:
- Industrial Employment (Standing Orders) Act, 1946
- Shops and Establishments Act (state-specific)
- Employees Provident Fund Act, 1952
- Payment of Gratuity Act, 1972
- Maternity Benefit Act, 1961
- Sexual Harassment of Women at Workplace Act, 2013 (POSH)

Include these sections:
1. DATE AND ADDRESSEE
2. OPENING — welcome and congratulate the candidate
3. DESIGNATION AND DEPARTMENT
4. COMPENSATION STRUCTURE
   - Gross CTC breakdown (basic, HRA, special allowance, etc.)
   - PF / EPF contribution
   - Gratuity eligibility
5. BENEFITS
   - Medical / health insurance
   - Leave entitlements (CL, SL, PL as per Shops Act)
   - Other perks
6. DATE OF JOINING AND WORK LOCATION
7. PROBATION PERIOD AND CONFIRMATION
8. NOTICE PERIOD
9. CODE OF CONDUCT AND POLICIES
   - POSH policy reference
   - Confidentiality obligation
10. CONDITIONS OF EMPLOYMENT
    - Background verification
    - Document submission
11. ACCEPTANCE INSTRUCTIONS
12. CLOSING
13. SIGNATURE BLOCK

Formatting:
- Warm professional tone for a letter
- Reference Indian statutes where appropriate
- CTC in INR (Indian Rupees)
- Plain text only, no markdown""",

    "freelancer_agreement": """You are an Indian contracts attorney specialising in IT and creative services law.
Draft a complete Freelancer / Independent Contractor Agreement under Indian law.

Legal framework references:
- Indian Contract Act, 1872
- Copyright Act, 1957 (Section 17 — work for hire)
- Information Technology Act, 2000
- Goods and Services Tax Act, 2017 (GST)
- Income Tax Act, 1961 (TDS under Section 194C / 194J)

Include these sections:
1. PARTIES AND RECITALS
2. SCOPE OF WORK
   2.1 Project Description
   2.2 Deliverables
   2.3 Timeline and Milestones
3. COMPENSATION AND PAYMENT
   3.1 Project Fee (in INR)
   3.2 Payment Schedule
   3.3 GST (if applicable — GSTIN of both parties)
   3.4 TDS Deduction (Section 194C or 194J as applicable)
   3.5 Late Payment Interest
4. REVISIONS AND CHANGE ORDERS
5. INTELLECTUAL PROPERTY
   5.1 Assignment of Copyright under Copyright Act, 1957
   5.2 Moral Rights Waiver
   5.3 Freelancer Portfolio Rights
6. CONFIDENTIALITY
7. INDEPENDENT CONTRACTOR STATUS
   7.1 No Employer-Employee Relationship
   7.2 Tax Responsibility
   7.3 No PF / ESI / Gratuity Obligation
8. CANCELLATION AND KILL FEE
9. LIMITATION OF LIABILITY
10. DISPUTE RESOLUTION
    - Arbitration under Arbitration and Conciliation Act, 1996
    - Seat of arbitration
11. GENERAL PROVISIONS
    - Governing Law (Indian law)
    - Stamp Duty
12. SIGNATURE BLOCK

Formatting:
- Section headings in ALL CAPS with number
- Sub-clauses numbered
- Formal legal language under Indian law
- All amounts in INR
- Plain text only, no markdown""",

    "service_agreement": """You are an Indian commercial contracts attorney.
Draft a complete Service Agreement governed by Indian law.

Legal framework references:
- Indian Contract Act, 1872
- Specific Relief Act, 1963
- Information Technology Act, 2000
- GST Act, 2017
- Arbitration and Conciliation Act, 1996

Include these sections:
1. PARTIES
2. DEFINITIONS
3. SCOPE OF SERVICES
   3.1 Services Description
   3.2 Service Standards and SLAs
   3.3 Change in Scope
4. TERM AND RENEWAL
5. FEES AND PAYMENT
   5.1 Service Fees (in INR)
   5.2 Payment Terms
   5.3 GST (GSTIN details)
   5.4 TDS Deduction
   5.5 Late Payment Interest
6. INTELLECTUAL PROPERTY
7. CONFIDENTIALITY AND DATA PROTECTION
   - Reference IT Act, 2000 and DPDP Act, 2023
8. REPRESENTATIONS AND WARRANTIES
9. LIMITATION OF LIABILITY
10. INDEMNIFICATION
11. TERMINATION
    11.1 Termination for Convenience
    11.2 Termination for Cause
    11.3 Effect of Termination
12. DISPUTE RESOLUTION
    - Arbitration under Arbitration and Conciliation Act, 1996
    - Governing Law: Indian law
    - Jurisdiction: courts of [governing state]
13. GENERAL PROVISIONS
14. SIGNATURE BLOCK

Formatting:
- Section headings in ALL CAPS with number
- Sub-clauses numbered
- Formal legal language under Indian law
- All amounts in INR
- Plain text only, no markdown""",

    "consulting_agreement": """You are an Indian commercial attorney specialising in professional services.
Draft a complete Consulting Agreement governed by Indian law.

Legal framework references:
- Indian Contract Act, 1872
- Copyright Act, 1957
- Income Tax Act, 1961 (TDS Section 194J)
- GST Act, 2017
- Arbitration and Conciliation Act, 1996

Include these sections:
1. PARTIES
2. SCOPE OF CONSULTING SERVICES
   2.1 Scope of Work
   2.2 Deliverables
   2.3 Consultant's Resources
3. TERM
4. COMPENSATION AND EXPENSES
   4.1 Consulting Fees (in INR)
   4.2 Expense Reimbursement
   4.3 GST (if applicable)
   4.4 TDS under Section 194J
   4.5 Invoicing and Payment Timeline
5. INTELLECTUAL PROPERTY
   5.1 Work Product Ownership
   5.2 Background IP
   5.3 Licence Grant
6. CONFIDENTIALITY
7. NON-COMPETE AND NON-SOLICITATION
   7.1 Non-Competition
   7.2 Non-Solicitation of Employees
   7.3 Non-Solicitation of Clients
8. INDEPENDENT CONTRACTOR STATUS
   - No PF / ESI / Gratuity obligation
9. REPRESENTATIONS AND WARRANTIES
10. LIMITATION OF LIABILITY
11. TERMINATION
12. DISPUTE RESOLUTION
    - Arbitration under Arbitration and Conciliation Act, 1996
    - Governing Law: Indian law
13. GENERAL PROVISIONS
14. SIGNATURE BLOCK

Formatting:
- Section headings in ALL CAPS with number
- Sub-clauses numbered
- Formal legal language under Indian law
- All amounts in INR
- Plain text only, no markdown""",

    "lease_agreement": """You are an Indian property law attorney specialising in residential and commercial leases.
Draft a complete Leave and Licence / Lease Agreement compliant with Indian property law.

Legal framework references:
- Transfer of Property Act, 1882
- Registration Act, 1908 (mandatory registration if term > 12 months)
- Indian Stamp Act, 1899 (state-specific stamp duty)
- Rent Control Act (state-specific)
- Maharashtra Rent Control Act / Delhi Rent Act (as applicable)

Include these sections:
1. PARTIES AND PROPERTY
   - Full names, addresses, Aadhaar/PAN references
   - Complete property description with area in sq.ft.
2. NATURE OF AGREEMENT
   - Leave and Licence (preferred) or Lease
   - Distinction from tenancy under Rent Control Act
3. LICENCE / LEASE TERM
   - Start date, end date, lock-in period
4. LICENCE FEE / RENT
   4.1 Monthly Amount (in INR)
   4.2 Due Date (e.g. 1st of each month)
   4.3 Mode of Payment (bank transfer, cheque, UPI)
   4.4 Annual Escalation clause (typically 5-10%)
5. REFUNDABLE SECURITY DEPOSIT
   4.1 Amount (in INR)
   4.2 Conditions for Deduction
   4.3 Return Timeline (typically 30-60 days after vacating)
6. MAINTENANCE AND SOCIETY CHARGES
7. UTILITIES AND SERVICES
8. PERMITTED USE
9. RESTRICTION ON SUBLETTING / ASSIGNMENT
10. ALTERATIONS AND REPAIRS
    10.1 Licensor / Landlord Obligations
    10.2 Licensee / Tenant Obligations
11. ENTRY BY LICENSOR / LANDLORD
12. TERMINATION AND NOTICE PERIOD
13. VACATION OF PREMISES AND MOVE-OUT
14. STAMP DUTY AND REGISTRATION
    - Applicable state stamp duty
    - Registration under Registration Act, 1908
15. GENERAL PROVISIONS
    - Governing Law and Jurisdiction
16. SCHEDULE A — Property Details
17. SIGNATURE BLOCK WITH WITNESSES

Formatting:
- Section headings in ALL CAPS with number
- Sub-clauses numbered
- Formal legal language under Indian property law
- All amounts in INR
- Reference state-specific laws where appropriate
- Plain text only, no markdown""",

    "employment_contract": """You are an Indian employment law attorney.
Draft a complete Employment Contract / Appointment Letter compliant with Indian labour law.

Legal framework references:
- Industrial Employment (Standing Orders) Act, 1946
- Employees Provident Fund and Miscellaneous Provisions Act, 1952
- Payment of Gratuity Act, 1972
- Payment of Bonus Act, 1965
- Maternity Benefit Act, 1961
- Sexual Harassment of Women at Workplace (Prevention, Prohibition and Redressal) Act, 2013
- Shops and Establishments Act (state-specific)
- Labour Codes (Code on Wages 2019, Code on Social Security 2020)
- Income Tax Act, 1961 (TDS on salary)
- Digital Personal Data Protection Act, 2023

Include these sections:
1. PARTIES
2. APPOINTMENT AND DESIGNATION
   2.1 Job Title and Grade
   2.2 Department and Reporting Structure
   2.3 Place of Work
3. DATE OF JOINING AND COMMENCEMENT
4. NATURE OF EMPLOYMENT
   - Permanent / Contract / Fixed-term
5. COMPENSATION AND BENEFITS
   5.1 Cost to Company (CTC) — annual in INR
   5.2 CTC Breakup (Basic, HRA, Special Allowance, LTA, etc.)
   5.3 Provident Fund (EPF) — 12% of Basic under EPF Act, 1952
   5.4 Gratuity — as per Payment of Gratuity Act, 1972
   5.5 Health and Medical Insurance
   5.6 Performance Bonus / Variable Pay
   5.7 Income Tax (TDS deduction)
6. WORKING HOURS
   - As per Shops and Establishments Act of [governing state]
7. LEAVE ENTITLEMENTS
   7.1 Earned / Privileged Leave
   7.2 Casual Leave
   7.3 Sick / Medical Leave
   7.4 Maternity / Paternity Leave
   7.5 National and Festival Holidays
8. PROBATIONARY PERIOD
   - Duration, confirmation process, extension
9. NOTICE PERIOD
   - During probation and post-confirmation
10. CODE OF CONDUCT AND POLICIES
    - Company policies, POSH compliance
    - Prevention of Sexual Harassment (POSH Act, 2013)
11. CONFIDENTIALITY AND NON-DISCLOSURE
12. INTELLECTUAL PROPERTY ASSIGNMENT
    - All work product vests in employer
    - Reference Copyright Act, 1957
13. NON-COMPETE AND NON-SOLICITATION
    - Reasonable restrictions under Section 27, Indian Contract Act
14. TERMINATION OF EMPLOYMENT
    14.1 Resignation by Employee
    14.2 Termination by Employer
    14.3 Termination for Cause (misconduct, etc.)
    14.4 Retirement
    14.5 Full and Final Settlement
15. POST-TERMINATION OBLIGATIONS
16. DATA PROTECTION
    - Digital Personal Data Protection Act, 2023
17. GRIEVANCE REDRESSAL
    - Internal Complaints Committee (POSH)
    - HR escalation process
18. GOVERNING LAW AND DISPUTE RESOLUTION
    - Governing law: Laws of India
    - Jurisdiction: courts of [governing state]
    - Arbitration under Arbitration and Conciliation Act, 1996
19. GENERAL PROVISIONS
20. SIGNATURE BLOCK

Formatting:
- Section headings in ALL CAPS with number
- Sub-clauses numbered
- Formal legal language under Indian law
- All amounts in INR
- Reference specific Indian statutes throughout
- Plain text only, no markdown""",
}


# ---------------------------------------------------------------------------
# Intent classification
# ---------------------------------------------------------------------------

_INTENT_SYSTEM = """You are a legal document classifier for Indian legal documents.
Read the description and return ONLY the matching slug.

Slugs:
nda                  - Non-Disclosure Agreement, confidentiality agreement
job_offer            - Job offer letter, appointment letter, offer of employment
freelancer_agreement - Freelancer contract, independent contractor agreement
service_agreement    - Service agreement, vendor agreement, SLA, AMC
consulting_agreement - Consulting agreement, advisory agreement, retainer
lease_agreement      - Lease, rent agreement, leave and licence, rental deed
employment_contract  - Employment contract, employment agreement, service contract

Return ONLY the slug. No explanation. If none match return unknown."""


async def classify_intent(user_query: str) -> str | None:
    client = _get_client()
    kwargs = _api_kwargs(max_tokens=20, use_json=False)
    kwargs["messages"] = [
        {"role": "system", "content": _INTENT_SYSTEM},
        {"role": "user",   "content": f"Classify:\n\"\"\"{user_query[:800]}\"\"\""},
    ]
    try:
        response = await client.chat.completions.create(**kwargs)
        raw      = (response.choices[0].message.content or "").strip().lower()
        raw      = re.sub(r'[^a-z_]', '', raw)
        if raw in SUPPORTED_DOCUMENT_TYPES:
            logger.info(f"[doc_gen] intent classified: '{raw}'")
            return raw
        resolved = resolve_document_type(raw)
        if resolved:
            return resolved
        logger.warning(f"[doc_gen] intent unknown: '{raw}'")
        return None
    except Exception as e:
        logger.error(f"[doc_gen] intent classification failed: {e}")
        return None


# ---------------------------------------------------------------------------
# Own JSON extractor
# ---------------------------------------------------------------------------

def _parse_json(raw: str) -> dict:
    text = raw.replace("```json", "").replace("```", "").strip()
    try:
        data = json.loads(text)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            data = json.loads(match.group())
            if isinstance(data, dict):
                return data
        except json.JSONDecodeError:
            pass
    logger.warning("[doc_gen] could not parse JSON")
    return {}


# ---------------------------------------------------------------------------
# Step 1: Extract fields
# ---------------------------------------------------------------------------

async def _extract_fields(document_type: str, user_query: str) -> dict[str, Any]:
    schema = _SCHEMAS[document_type]
    system = _EXTRACTION_PROMPTS[document_type]

    client = _get_client()
    kwargs = _api_kwargs(max_tokens=4096, use_json=True)
    kwargs["messages"] = [
        {"role": "system", "content": system},
        {"role": "user",   "content": (
            f"Extract all fields from this description as a JSON object.\n\n"
            f"Description:\n\"\"\"{user_query}\"\"\""
        )},
    ]

    response = await client.chat.completions.create(**kwargs)
    raw      = response.choices[0].message.content or "{}"
    logger.info(
        f"[doc_gen] field extraction "
        f"tokens={response.usage.prompt_tokens}in/{response.usage.completion_tokens}out "
        f"finish={response.choices[0].finish_reason}"
    )

    fields = _parse_json(raw)

    for field in schema["required"]:
        if field not in fields or not fields[field]:
            fields[field] = "Not Specified"

    found = {k: v for k, v in fields.items() if v != "Not Specified"}
    logger.info(f"[doc_gen] extracted {len(found)} non-empty fields: {list(found.keys())}")

    return fields


# ---------------------------------------------------------------------------
# Step 2: Validate required fields
# ---------------------------------------------------------------------------

def _get_missing_fields(document_type: str, fields: dict) -> list[str]:
    return [
        f for f in _SCHEMAS[document_type]["required"]
        if not fields.get(f) or fields.get(f) == "Not Specified"
    ]


# ---------------------------------------------------------------------------
# Step 3: Generate document text
# ---------------------------------------------------------------------------

async def _generate_document_text(
    document_type: str,
    fields: dict[str, Any],
    user_query: str,
) -> tuple[str, int, int]:
    doc_name   = SUPPORTED_DOCUMENT_TYPES[document_type]
    gen_prompt = _GENERATION_PROMPTS[document_type]

    field_lines = "\n".join(
        f"  {k.replace('_', ' ').title()}: {v}"
        for k, v in fields.items()
        if v and v != "Not Specified"
    )

    system = gen_prompt

    user = f"""Use these details to generate the complete {doc_name} under Indian law:

{field_lines if field_lines else "(Use standard Indian law template values)"}

Additional context from user:
\"\"\"{user_query}\"\"\"

Write the complete {doc_name} now. Start directly with the document title."""

    client = _get_client()
    kwargs = _api_kwargs(max_tokens=32000, use_json=False)
    kwargs["messages"] = [
        {"role": "system", "content": system},
        {"role": "user",   "content": user},
    ]

    response      = await client.chat.completions.create(**kwargs)
    document_text = response.choices[0].message.content or ""
    in_tok        = response.usage.prompt_tokens
    out_tok       = response.usage.completion_tokens
    finish        = response.choices[0].finish_reason

    logger.info(
        f"[doc_gen] document generated — type={document_type} "
        f"length={len(document_text)} chars "
        f"tokens={in_tok}in/{out_tok}out finish={finish}"
    )

    if not document_text.strip():
        logger.error(f"[doc_gen] empty response — finish={finish}")

    return document_text, in_tok, out_tok


# ---------------------------------------------------------------------------
# Step 4: Render DOCX
# ---------------------------------------------------------------------------

def _render_docx(document_text: str, document_name: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime

    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    BLUE  = RGBColor(0x2c, 0x4a, 0x8c)
    DARK  = RGBColor(0x1a, 0x1a, 0x1a)
    MUTED = RGBColor(0x55, 0x55, 0x55)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    def _set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _remove_table_borders(table):
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "none")
            tblBorders.append(b)
        tblPr.append(tblBorders)

    # Header banner
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell  = table.cell(0, 0)
    _set_cell_bg(cell, "1a2744")
    _remove_table_borders(table)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run(document_name.upper())
    run.bold           = True
    run.font.size      = Pt(20)
    run.font.color.rgb = WHITE

    # India jurisdiction tag
    para2 = cell.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2  = para2.add_run("Governed by the Laws of India")
    run2.font.size      = Pt(9)
    run2.font.color.rgb = RGBColor(0xCC, 0xD6, 0xF0)

    doc.add_paragraph()

    dp  = doc.add_paragraph()
    dr  = dp.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    dr.font.size      = Pt(9)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    dp.paragraph_format.space_after = Pt(12)

    lines        = document_text.split("\n")
    in_signature = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(3)
            continue

        if re.match(r'^\d+\.\s+SIGNATURE', stripped.upper()):
            in_signature = True

        is_section = (
            re.match(r'^\d+\.\s+[A-Z][A-Z\s,/&().-]{3,}[:.]?\s*$', stripped)
            or re.match(r'^[A-Z][A-Z\s,/&().-]{3,}[:.]?\s*$', stripped)
            or (stripped.isupper() and 4 < len(stripped) < 80)
        )

        if is_section and not in_signature:
            p    = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            top  = OxmlElement("w:top")
            top.set(qn("w:val"),   "single")
            top.set(qn("w:sz"),    "4")
            top.set(qn("w:space"), "1")
            top.set(qn("w:color"), "CCCCCC")
            pBdr.append(top)
            pPr.append(pBdr)
            run = p.add_run(stripped)
            run.bold           = True
            run.font.size      = Pt(11)
            run.font.color.rgb = BLUE
            continue

        if re.match(r'^\d+\.\d+\.?\s+', stripped) and not in_signature:
            p   = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.font.size      = Pt(10)
            run.font.color.rgb = DARK
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue

        if in_signature and re.match(r'^(Name|Title|Date|Company|Designation|Signature|Witness|Place|PAN|Aadhaar)\s*:', stripped, re.I):
            label, _, rest = stripped.partition(":")
            lp  = doc.add_paragraph()
            lr  = lp.add_run(label.strip() + ":")
            lr.bold            = True
            lr.font.size       = Pt(9)
            lr.font.color.rgb  = MUTED
            lp.paragraph_format.space_after = Pt(2)
            vp  = doc.add_paragraph()
            vr  = vp.add_run(rest.strip() if rest.strip() else "________________________")
            vr.font.size      = Pt(10)
            vr.font.color.rgb = DARK
            vp.paragraph_format.space_after = Pt(12)
            continue

        p   = doc.add_paragraph(style="Normal")
        run = p.add_run(stripped)
        run.font.size      = Pt(10)
        run.font.color.rgb = DARK
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)

    # Footer
    for section in doc.sections:
        footer      = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"{document_name}  |  Governed by Laws of India  |  Confidential  |  Page "
        )
        run.font.size      = Pt(8)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        fldChar1              = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText             = OxmlElement("w:instrText")
        instrText.text        = "PAGE"
        fldChar2              = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run2 = footer_para.add_run()
        run2.font.size      = Pt(8)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)

    buffer     = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()

    logger.info(f"[doc_gen] DOCX rendered — {len(docx_bytes):,} bytes")
    return docx_bytes


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def generate_document(
    document_type: str | None,
    user_query: str,
) -> dict:
    total_in_tok  = 0
    total_out_tok = 0

    if document_type:
        doc_type = resolve_document_type(document_type)
        if not doc_type:
            return {
                "status":        "unknown_type",
                "document_type": document_type,
                "document_name": "",
                "fields":        {},
                "missing_fields": [],
                "document":      "",
                "docx_bytes":    b"",
                "word_count":    0,
                "input_tokens":  0,
                "output_tokens": 0,
                "message": (
                    f"'{document_type}' is not supported. "
                    f"Supported: {', '.join(SUPPORTED_DOCUMENT_TYPES.keys())}"
                ),
            }
    else:
        logger.info("[doc_gen] document_type not provided — classifying intent...")
        doc_type = await classify_intent(user_query)
        if not doc_type:
            return {
                "status":        "unknown_type",
                "document_type": None,
                "document_name": "",
                "fields":        {},
                "missing_fields": [],
                "document":      "",
                "docx_bytes":    b"",
                "word_count":    0,
                "input_tokens":  total_in_tok,
                "output_tokens": total_out_tok,
                "message": (
                    "Could not determine document type. "
                    "Please specify document_type. "
                    f"Supported: {', '.join(SUPPORTED_DOCUMENT_TYPES.keys())}"
                ),
            }

    doc_name = SUPPORTED_DOCUMENT_TYPES[doc_type]
    logger.info(f"[doc_gen] generating '{doc_name}' ({len(user_query)} chars)")

    fields  = await _extract_fields(doc_type, user_query)
    missing = _get_missing_fields(doc_type, fields)
    if missing:
        logger.warning(f"[doc_gen] missing: {missing}")

    document_text, in_tok, out_tok = await _generate_document_text(
        doc_type, fields, user_query
    )
    total_in_tok  += in_tok
    total_out_tok += out_tok

    docx_bytes = _render_docx(document_text, doc_name)

    return {
        "status":         "missing_fields" if missing else "success",
        "document_type":  doc_type,
        "document_name":  doc_name,
        "fields":         fields,
        "missing_fields": missing,
        "document":       document_text,
        "docx_bytes":     docx_bytes,
        "word_count":     len(document_text.split()),
        "input_tokens":   total_in_tok,
        "output_tokens":  total_out_tok,
    }
import re
import fitz
import time
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Structure detection helpers
# ---------------------------------------------------------------------------

def _detect_block_type(text: str, font_size: float = 12, is_bold: bool = False) -> str:
    """
    Classify a text block as heading, subheading, list item, or paragraph.
    """
    text = text.strip()
    if not text:
        return "empty"

    # Heading: short, bold or large font, no full stop at end
    if (is_bold or font_size >= 16) and len(text) < 120 and not text.endswith("."):
        return "heading1"

    if font_size >= 13 and len(text) < 150 and not text.endswith("."):
        return "heading2"

    # List item: starts with bullet, dash, number
    if re.match(r'^[\•\-\*\u2022\u2023\u25E6]\s+', text):
        return "list_bullet"
    if re.match(r'^\d+[\.\)]\s+', text):
        return "list_number"

    return "paragraph"


def _extract_blocks_native(page: fitz.Page) -> list[dict]:
    """
    Extract text blocks from a native PDF page with font metadata.
    Returns list of {text, font_size, is_bold, block_type}
    """
    blocks = []
    try:
        raw_blocks = page.get_text("dict")["blocks"]
        for block in raw_blocks:
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                line_text  = ""
                font_size  = 12
                is_bold    = False
                for span in line.get("spans", []):
                    line_text += span.get("text", "")
                    font_size  = span.get("size", 12)
                    flags      = span.get("flags", 0)
                    is_bold    = bool(flags & 2**4)  # bit 4 = bold

                line_text = line_text.strip()
                if not line_text:
                    continue

                blocks.append({
                    "text":       line_text,
                    "font_size":  font_size,
                    "is_bold":    is_bold,
                    "block_type": _detect_block_type(line_text, font_size, is_bold),
                })
    except Exception as e:
        logger.warning(f"_extract_blocks_native: {e}")

    return blocks


def _extract_blocks_ocr(ocr_text: str) -> list[dict]:
    """
    Convert flat OCR text into structured blocks.
    Without font metadata, use heuristics on line length and content.
    """
    blocks = []
    for line in ocr_text.splitlines():
        line = line.strip()
        if not line:
            continue
        blocks.append({
            "text":       line,
            "font_size":  12,
            "is_bold":    False,
            "block_type": _detect_block_type(line),
        })
    return blocks


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def _build_docx(all_blocks: list[list[dict]], source_filename: str) -> Document:
    """
    Build a DOCX Document from extracted blocks.
    all_blocks is a list of pages, each page is a list of block dicts.
    """
    doc = Document()

    # Document title from filename
    title_para = doc.add_heading(
        source_filename.replace(".pdf", "").replace("_", " ").title(), level=0
    )
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for page_num, page_blocks in enumerate(all_blocks, 1):
        # Page separator (except before first page)
        if page_num > 1:
            doc.add_paragraph()  # spacing between pages

        for block in page_blocks:
            btype = block["block_type"]
            text  = block["text"]

            if btype == "empty":
                continue

            elif btype == "heading1":
                doc.add_heading(text, level=1)

            elif btype == "heading2":
                doc.add_heading(text, level=2)

            elif btype == "list_bullet":
                # Strip leading bullet character
                clean = re.sub(r'^[\•\-\*\u2022\u2023\u25E6]\s+', '', text)
                doc.add_paragraph(clean, style="List Bullet")

            elif btype == "list_number":
                clean = re.sub(r'^\d+[\.\)]\s+', '', text)
                doc.add_paragraph(clean, style="List Number")

            else:  # paragraph
                para = doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def pdf_to_docx(file_path: str, filename: str) -> Document:
    """
    Convert a PDF file to a DOCX Document object.

    Uses:
      - PyMuPDF native extraction with font metadata for text PDFs
      - Falls back to flat OCR text (passed in) for scanned PDFs

    Returns a python-docx Document ready to be saved.
    """
    t0 = time.perf_counter()
    logger.info(f"[pdf_to_docx] converting '{filename}'")

    try:
        doc_fitz = fitz.open(file_path)
    except Exception as e:
        raise ValueError(f"Could not open PDF: {e}")

    all_blocks = []

    for page_num in range(len(doc_fitz)):
        page   = doc_fitz[page_num]
        blocks = _extract_blocks_native(page)

        # If page has no native text, mark for OCR
        if not blocks:
            logger.info(f"[pdf_to_docx] page {page_num+1}: no native text — needs OCR")
            # OCR is handled in the endpoint using existing PaddleOCR-VL
            # Placeholder so page count stays accurate
            blocks = [{"text": f"[Page {page_num+1}: scanned content]",
                       "font_size": 12, "is_bold": False, "block_type": "paragraph"}]

        all_blocks.append(blocks)

    doc_fitz.close()

    docx = _build_docx(all_blocks, filename)
    logger.info(f"[pdf_to_docx] done ({time.perf_counter()-t0:.2f}s) — {len(all_blocks)} page(s)")
    return docx